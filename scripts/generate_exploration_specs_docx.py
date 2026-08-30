#!/usr/bin/env python3
"""Generate the exploration SFD/STD in the editorial format used by Project Sources."""

from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent
TEMPLATE = WORKSPACE / "project_sources" / "03-SFD-Module-paiement-marchand-1-.docx"
SOURCE_DIR = ROOT / "docs" / "design" / "Sources"
DELIVERABLE_DIR = WORKSPACE / "deliverables"

INK = RGBColor(28, 32, 39)
SLATE = RGBColor(82, 92, 108)
MUTED = RGBColor(137, 151, 170)
LIGHT = "EEF1F5"
LINE = "D8DEE7"


DOCS = [
    {
        "kind": "SFD · SPÉCIFICATIONS FONCTIONNELLES DÉTAILLÉES",
        "title": "Exploration vivante",
        "title2": "Acteurs & déclencheurs",
        "subtitle": "PNJ autonomes, poursuite ennemie, contact de combat et exceptions scénarisées.",
        "reference": "LEDS-SFD-EXP-001",
        "status": "Référence d'implémentation",
        "source": ROOT / "docs" / "design" / "sfd-exploration-acteurs-et-declencheurs.md",
        "output": "LEDS-SFD-EXP-001_Exploration_acteurs_et_declencheurs_v1.0.docx",
    },
    {
        "kind": "STD · SPÉCIFICATIONS TECHNIQUES DÉTAILLÉES",
        "title": "Moteur d’acteurs",
        "title2": "d’exploration",
        "subtitle": "Contrats domaine, API, persistance et animation du nouveau modèle d'exploration.",
        "reference": "LEDS-STD-EXP-001",
        "status": "Référence d'implémentation",
        "source": ROOT / "docs" / "design" / "std-exploration-acteurs-et-declencheurs.md",
        "output": "LEDS-STD-EXP-001_Moteur_acteurs_exploration_v1.0.docx",
    },
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_table_borders(table, color: str = LINE, size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_table_fixed_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(index, len(widths) - 1)])


def set_run(run, *, size: float = 10.5, bold: bool = False, italic: bool = False,
            color: RGBColor = INK, font: str = "Georgia") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def replace_paragraph(paragraph, text: str) -> None:
    runs = paragraph.runs
    if not runs:
        runs = [paragraph.add_run()]
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def prepare_template(meta: dict) -> Document:
    doc = Document(TEMPLATE)
    body = doc._element.body
    # Keep the template's cover table, its first-section break, and the final section properties.
    for child in list(body)[2:-1]:
        body.remove(child)

    cover = doc.tables[0].cell(0, 0)
    replace_paragraph(cover.paragraphs[0], "L’ÉPOPÉE DES SILENCES\tCONFIDENTIEL — DOCUMENT DE TRAVAIL")
    replace_paragraph(cover.paragraphs[2], meta["kind"])
    replace_paragraph(cover.paragraphs[3], meta["title"])
    replace_paragraph(cover.paragraphs[4], meta["title2"])
    replace_paragraph(cover.paragraphs[5], meta["subtitle"])

    metadata = cover.tables[0]
    values = [meta["reference"], "1.0 — 25/08/2026", "Équipe L’Épopée des Silences", meta["status"]]
    for index, value in enumerate(values):
        replace_paragraph(metadata.cell(1, index).paragraphs[0], value)

    return doc


def configure_body(doc: Document, running_label: str) -> None:
    section = doc.sections[-1]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    replace_paragraph(p, f"L’ÉPOPÉE DES SILENCES    ·    {running_label}")
    for run in p.runs:
        set_run(run, size=7.6, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("DOCUMENT DE RÉFÉRENCE    ·    ")
    set_run(run, size=7.2, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Georgia"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15


def add_kicker(doc: Document, number: str, label: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{number}    {label.upper()}")
    set_run(r, size=8.2, bold=True, color=MUTED)
    r.font.all_caps = True


def add_heading(doc: Document, text: str, level: int, counter: list[int]) -> None:
    if level == 2:
        counter[0] += 1
        counter[1] = 0
        add_kicker(doc, f"{counter[0]:02d}", "SECTION")
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_after = Pt(9)
        r = p.add_run(text)
        set_run(r, size=19, bold=True, color=INK)
    else:
        counter[1] += 1
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(text)
        set_run(r, size=13.2, bold=True, color=SLATE)


def add_rich_paragraph(doc: Document, text: str, *, bullet: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22 if bullet else 0)
    p.paragraph_format.first_line_indent = Inches(-0.14 if bullet else 0)
    p.paragraph_format.space_after = Pt(5)
    if bullet:
        lead = p.add_run("—  ")
        set_run(lead, size=10.4, bold=True, color=SLATE)
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        code = part.startswith("`") and part.endswith("`")
        bold = part.startswith("**") and part.endswith("**")
        value = part[1:-1] if code else part[2:-2] if bold else part
        run = p.add_run(value)
        set_run(run, size=9.6 if code else 10.5, bold=bold, color=INK,
                font="Consolas" if code else "Georgia")


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed_width(table, [9360])
    set_table_borders(table, color="D5DAE2", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F4F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("\n".join(lines))
    set_run(r, size=8.2, color=SLATE, font="Consolas")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    widths = [9360 // cols] * cols
    widths[-1] += 9360 - sum(widths)
    set_table_fixed_width(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for r_index, values in enumerate(rows):
        for c_index, value in enumerate(values):
            cell = table.cell(r_index, c_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_index == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(value.strip())
            set_run(run, size=8.8, bold=r_index == 0, color=SLATE if r_index == 0 else INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def sections_from_markdown(text: str) -> list[tuple[str, str]]:
    result = []
    for line in text.splitlines():
        if line.startswith("## "):
            title = re.sub(r"^\d+\.\s*", "", line[3:].strip())
            result.append((f"{len(result) + 1:02d}", title))
    return result


def add_summary(doc: Document, text: str) -> None:
    add_kicker(doc, "—", "SOMMAIRE")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Sommaire")
    set_run(r, size=22, bold=True)
    sections = sections_from_markdown(text)
    table = doc.add_table(rows=len(sections), cols=2)
    set_table_fixed_width(table, [900, 8460])
    set_table_borders(table, color="E0E5EC", size="4")
    for i, (number, title) in enumerate(sections):
        for c, value in enumerate((number, title)):
            cell = table.cell(i, c)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(value)
            set_run(run, size=10.2, bold=c == 1, color=MUTED if c == 0 else INK)
    doc.add_page_break()


def add_markdown_body(doc: Document, text: str) -> None:
    lines = text.splitlines()
    counter = [0, 0]
    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        line = lines[index].rstrip()
        if line.startswith("# ") or line.startswith("Statut :") or line.startswith("Référence :") or line.startswith("Version :"):
            index += 1
            continue
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if line.startswith("| ") and index + 1 < len(lines) and re.match(r"^\|[-: |]+\|$", lines[index + 1]):
            rows = [[part.strip() for part in line.strip("|").split("|")]]
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                rows.append([part.strip() for part in lines[index].strip("|").split("|")])
                index += 1
            add_markdown_table(doc, rows)
            continue
        if line.startswith("## "):
            heading = re.sub(r"^\d+\.\s*", "", line[3:].strip()).strip("`")
            add_heading(doc, heading, 2, counter)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip().strip("`"), 3, counter)
        elif line.startswith("- "):
            add_rich_paragraph(doc, line[2:].strip(), bullet=True)
        elif re.match(r"^\d+\.\s", line):
            add_rich_paragraph(doc, line, bullet=True)
        elif line.strip():
            add_rich_paragraph(doc, line.strip())
        index += 1


def generate(meta: dict) -> Path:
    text = meta["source"].read_text(encoding="utf-8")
    doc = prepare_template(meta)
    configure_body(doc, meta["reference"])
    add_summary(doc, text)
    add_markdown_body(doc, text)

    SOURCE_DIR.mkdir(parents=True, exist_ok=True)
    DELIVERABLE_DIR.mkdir(parents=True, exist_ok=True)
    output = SOURCE_DIR / meta["output"]
    doc.save(output)
    shutil.copy2(output, DELIVERABLE_DIR / output.name)
    return output


if __name__ == "__main__":
    for document in DOCS:
        print(generate(document))

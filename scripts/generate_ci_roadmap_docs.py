#!/usr/bin/env python3
"""Generate the CI/CD reference and beta roadmap DOCX files."""

from __future__ import annotations

import os
import re
import shutil
from pathlib import Path

os.environ.setdefault("MPLCONFIGDIR", "/tmp/leds-matplotlib")

import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.parser import parse_xml
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.parts.numbering import NumberingPart
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent
TEMPLATE = WORKSPACE / "project_sources" / "03-SFD-Module-paiement-marchand-1-.docx"
SOURCE_DIR = ROOT / "docs" / "design" / "Sources"
DELIVERABLE_DIR = WORKSPACE / "deliverables"
ASSET_DIR = Path("/tmp/leds-ci-roadmap-assets")

INK = RGBColor(28, 32, 39)
SLATE = RGBColor(82, 92, 108)
MUTED = RGBColor(137, 151, 170)
NAVY_HEX = "1C2027"
SLATE_HEX = "525C6C"
MUTED_HEX = "8997AA"
LIGHT = "EEF1F5"
LIGHTER = "F7F8FA"
LINE = "D8DEE7"
WHITE = "FFFFFF"
CONTENT_WIDTH = 9360


DOCS = [
    {
        "kind": "STD · SPÉCIFICATIONS TECHNIQUES DÉTAILLÉES",
        "title": "Plateforme CI/CD",
        "title2": "Qualité & livraison",
        "subtitle": "Tests, couverture, sécurité de la supply chain et déploiement reproductible.",
        "reference": "LEDS-STD-CICD-001",
        "status": "Référence de mise en œuvre",
        "source": ROOT / "docs" / "design" / "std-ci-cd.md",
        "output": "LEDS-STD-CICD-001_Plateforme_CI_CD_v1.0.docx",
    },
    {
        "kind": "PLAN · ROADMAP DE DÉVELOPPEMENT",
        "title": "Roadmap vers",
        "title2": "la bêta jouable",
        "subtitle": "CI/CD, comptes utilisateurs, tranche verticale, contenu et durcissement.",
        "reference": "LEDS-PLAN-BETA-001",
        "status": "Plan directeur",
        "source": ROOT / "docs" / "design" / "plan-roadmap-beta.md",
        "output": "LEDS-PLAN-BETA-001_Roadmap_vers_beta_v1.0.docx",
    },
]


def set_run(run, *, size: float = 10.2, bold: bool = False, italic: bool = False,
            color: RGBColor = INK, font: str = "Georgia") -> None:
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def replace_paragraph(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run()
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 90, start: int = 120,
                     bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_table_borders(table, color: str = LINE, size: str = "5") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_table_fixed_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(index, len(widths) - 1)])
            set_cell_margins(cell)


def preferred_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    scores = []
    for col in range(cols):
        values = [re.sub(r"[*`]", "", row[col]) for row in rows]
        maximum = max(len(value) for value in values)
        average = sum(len(value) for value in values) / len(values)
        scores.append(max(9.0, min(45.0, 0.65 * maximum + 0.35 * average)))
    if cols >= 3:
        scores[0] = max(scores[0], 15.0)
    widths = [max(820, int(CONTENT_WIDTH * score / sum(scores))) for score in scores]
    delta = CONTENT_WIDTH - sum(widths)
    widths[scores.index(max(scores))] += delta
    return widths


def prepare_template(meta: dict) -> Document:
    doc = Document(TEMPLATE)
    body = doc._element.body
    for child in list(body)[2:-1]:
        body.remove(child)

    cover = doc.tables[0].cell(0, 0)
    replace_paragraph(cover.paragraphs[0], "L’ÉPOPÉE DES SILENCES\tCONFIDENTIEL — DOCUMENT DE TRAVAIL")
    replace_paragraph(cover.paragraphs[2], meta["kind"])
    replace_paragraph(cover.paragraphs[3], meta["title"])
    replace_paragraph(cover.paragraphs[4], meta["title2"])
    replace_paragraph(cover.paragraphs[5], meta["subtitle"])
    metadata = cover.tables[0]
    values = [meta["reference"], "1.0 — 25/08/2026", "Équipe L’Épopée des Silences", meta["status"]]
    for index, value in enumerate(values):
        replace_paragraph(metadata.cell(1, index).paragraphs[0], value)
    return doc


def configure_body(doc: Document, running_label: str) -> None:
    section = doc.sections[-1]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    replace_paragraph(p, f"L’ÉPOPÉE DES SILENCES    ·    {running_label}")
    for run in p.runs:
        set_run(run, size=7.6, bold=True, color=MUTED)

    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    replace_paragraph(p, "")
    run = p.add_run("DOCUMENT DE RÉFÉRENCE    ·    ")
    set_run(run, size=7.2, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)

    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for name, outline_level in (("Heading 1", 0), ("Heading 2", 1)):
        style = doc.styles[name] if name in doc.styles else doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        style.next_paragraph_style = normal
        ppr = style._element.get_or_add_pPr()
        outline = ppr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            ppr.append(outline)
        outline.set(qn("w:val"), str(outline_level))


def add_kicker(doc: Document, number: str, label: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{number}    {label.upper()}")
    set_run(r, size=8.2, bold=True, color=MUTED)
    r.font.all_caps = True


def add_heading(doc: Document, text: str, level: int, counter: list[int]) -> None:
    if level == 2:
        counter[0] += 1
        counter[1] = 0
        add_kicker(doc, f"{counter[0]:02d}", "SECTION")
        p = doc.add_paragraph(style="Heading 1")
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_after = Pt(9)
        r = p.add_run(text)
        set_run(r, size=19, bold=True, color=INK)
    else:
        counter[1] += 1
        p = doc.add_paragraph(style="Heading 2")
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(text)
        set_run(r, size=13.2, bold=True, color=SLATE)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), SLATE_HEX)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Georgia")
    rfonts.set(qn("w:hAnsi"), "Georgia")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "19")
    rpr.extend([rfonts, color, underline, size])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([rpr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline_runs(paragraph, text: str, base_size: float = 10.2) -> None:
    pattern = r"(https?://\S+|`[^`]+`|\*\*[^*]+\*\*)"
    for part in re.split(pattern, text):
        if not part:
            continue
        if part.startswith("http://") or part.startswith("https://"):
            trimmed = part.rstrip(".,;)")
            suffix = part[len(trimmed):]
            add_hyperlink(paragraph, trimmed, trimmed)
            if suffix:
                run = paragraph.add_run(suffix)
                set_run(run, size=base_size)
        else:
            code = part.startswith("`") and part.endswith("`")
            bold = part.startswith("**") and part.endswith("**")
            value = part[1:-1] if code else part[2:-2] if bold else part
            run = paragraph.add_run(value)
            set_run(run, size=9.2 if code else base_size, bold=bold,
                    font="Consolas" if code else "Georgia")


def create_numbering(doc: Document, ordered: bool) -> int:
    try:
        numbering_part = doc.part.part_related_by(RT.NUMBERING)
    except KeyError:
        numbering_root = parse_xml(
            '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        numbering_part = NumberingPart(
            PackURI("/word/numbering.xml"),
            CT.WML_NUMBERING,
            numbering_root,
            doc.part.package,
        )
        doc.part.relate_to(numbering_part, RT.NUMBERING)
    numbering = numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "—")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "260")
    ppr.extend([tabs, ind])
    level.extend([start, num_fmt, lvl_text, lvl_jc, ppr])
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    ppr.append(num_pr)


def add_rich_paragraph(doc: Document, text: str, *, list_num_id: int | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_together = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if list_num_id is not None:
        apply_numbering(p, list_num_id)
    add_inline_runs(p, text)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed_width(table, [CONTENT_WIDTH])
    set_table_borders(table, color="D5DAE2", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F4F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("\n".join(lines))
    set_run(r, size=8.0, color=SLATE, font="Consolas")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def clean_md(value: str) -> str:
    return value.replace("**", "").replace("`", "").strip()


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    cols = len(rows[0])
    rows = [[clean_md(value) for value in row] for row in rows]
    table = doc.add_table(rows=len(rows), cols=cols)
    widths = preferred_widths(rows)
    if rows[0] == ["Version", "Date", "Évolution", "Statut"]:
        widths = [1200, 1600, 4960, 1600]
    set_table_fixed_width(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, LIGHT)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHTER)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.08
            alignment = WD_ALIGN_PARAGRAPH.CENTER if (len(value) <= 18 and col_index != 1) else WD_ALIGN_PARAGRAPH.LEFT
            p.alignment = alignment
            r = p.add_run(value)
            set_run(r, size=8.2, bold=row_index == 0, color=SLATE if row_index == 0 else INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def sections_from_markdown(text: str) -> list[tuple[str, str]]:
    sections = []
    for line in text.splitlines():
        if line.startswith("## "):
            title = re.sub(r"^\d+\.\s*", "", line[3:].strip())
            sections.append((f"{len(sections) + 1:02d}", title))
    return sections


def add_summary(doc: Document, text: str) -> None:
    add_kicker(doc, "—", "SOMMAIRE")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Sommaire")
    set_run(r, size=22, bold=True)
    sections = sections_from_markdown(text)
    table = doc.add_table(rows=len(sections), cols=2)
    set_table_fixed_width(table, [900, 8460])
    set_table_borders(table, color="E0E5EC", size="4")
    for row_index, (number, title) in enumerate(sections):
        for col_index, value in enumerate((number, title)):
            cell = table.cell(row_index, col_index)
            set_cell_shading(cell, LIGHTER if row_index % 2 else WHITE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(value)
            set_run(r, size=9.1, bold=col_index == 1, color=MUTED if col_index == 0 else INK)
    doc.add_page_break()


def add_figure(doc: Document, filename: str, caption: str) -> None:
    path = ASSET_DIR / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    picture = run.add_picture(str(path), width=Inches(6.35))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_before = Pt(2)
    c.paragraph_format.space_after = Pt(8)
    r = c.add_run(caption)
    set_run(r, size=8.2, italic=True, color=MUTED)


def add_markdown_body(doc: Document, text: str) -> None:
    lines = text.splitlines()
    counter = [0, 0]
    bullet_num_id = create_numbering(doc, ordered=False)
    ordered_num_id: int | None = None
    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        line = lines[index].rstrip()
        is_ordered_line = bool(re.match(r"^\d+\.\s", line))
        if not is_ordered_line:
            ordered_num_id = None
        if line.startswith("# ") or line.startswith(("Statut :", "Référence :", "Version :")):
            index += 1
            continue
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
            in_code = not in_code
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        figure = re.match(r"^\[\[FIGURE:([^|]+)\|(.+)\]\]$", line)
        if figure:
            add_figure(doc, figure.group(1), figure.group(2))
            index += 1
            continue
        if line == "[[PAGEBREAK]]":
            doc.add_page_break()
            index += 1
            continue
        if line.startswith("| ") and index + 1 < len(lines) and re.match(r"^\|[-: |]+\|$", lines[index + 1]):
            rows = [[part.strip() for part in line.strip("|").split("|")]]
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                rows.append([part.strip() for part in lines[index].strip("|").split("|")])
                index += 1
            add_markdown_table(doc, rows)
            continue
        if line.startswith("## "):
            add_heading(doc, re.sub(r"^\d+\.\s*", "", line[3:].strip()), 2, counter)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3, counter)
        elif line.startswith("- "):
            add_rich_paragraph(doc, line[2:].strip(), list_num_id=bullet_num_id)
        elif is_ordered_line:
            if ordered_num_id is None:
                ordered_num_id = create_numbering(doc, ordered=True)
            add_rich_paragraph(doc, re.sub(r"^\d+\.\s*", "", line), list_num_id=ordered_num_id)
        elif line.strip():
            add_rich_paragraph(doc, line.strip())
        index += 1


def box(ax, x: float, y: float, w: float, h: float, title: str, detail: str,
        fill: str = "#EEF1F5", edge: str = "#8997AA") -> None:
    patch = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.018,rounding_size=0.025",
                           linewidth=1.4, edgecolor=edge, facecolor=fill)
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h * 0.64, title, ha="center", va="center",
            fontsize=11, fontweight="bold", color="#1C2027")
    ax.text(x + w / 2, y + h * 0.30, detail, ha="center", va="center",
            fontsize=8.6, color="#525C6C", wrap=True)


def arrow(ax, start: tuple[float, float], end: tuple[float, float]) -> None:
    ax.add_patch(FancyArrowPatch(start, end, arrowstyle="-|>", mutation_scale=13,
                                linewidth=1.2, color="#8997AA"))


def save_figure(fig, name: str) -> None:
    path = ASSET_DIR / name
    fig.savefig(path, dpi=190, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def generate_diagrams() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    fig, ax = plt.subplots(figsize=(10.5, 6.0))
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    box(ax, .06, .78, .88, .14, "Pull request / merge / release", "GitHub · branches protégées · concurrence contrôlée", "#EAEFF7", "#7188A8")
    box(ax, .06, .56, .26, .14, "Détection monorepo", "Paths + dépendances partagées")
    box(ax, .37, .56, .26, .14, "Matrices CI", ".NET · Vitest · Playwright")
    box(ax, .68, .56, .26, .14, "Sécurité", "CodeQL · dépendances · secrets")
    arrow(ax, (.5, .78), (.19, .70)); arrow(ax, (.5, .78), (.5, .70)); arrow(ax, (.5, .78), (.81, .70))
    box(ax, .12, .33, .34, .14, "Gates de qualité", "0 test rouge · couverture ≥ 80 %")
    box(ax, .54, .33, .34, .14, "Build OCI unique", "4 images · 3 bundles EF")
    arrow(ax, (.19, .56), (.29, .47)); arrow(ax, (.50, .56), (.29, .47)); arrow(ax, (.81, .56), (.71, .47))
    arrow(ax, (.46, .40), (.54, .40))
    box(ax, .06, .08, .26, .14, "Intégration", "Automatique sur develop", "#E8F2EE", "#6C9A87")
    box(ax, .37, .08, .26, .14, "Préproduction", "RC + approbation", "#F4F0E5", "#AA9360")
    box(ax, .68, .08, .26, .14, "Bêta", "Release + approbation", "#EAEFF7", "#7188A8")
    arrow(ax, (.71, .33), (.19, .22)); arrow(ax, (.32, .15), (.37, .15)); arrow(ax, (.63, .15), (.68, .15))
    save_figure(fig, "ci-architecture.png")

    fig, ax = plt.subplots(figsize=(10.5, 6.0))
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    stages = [
        ("1 · Préflight", "Digests · attestations · verrou"),
        ("2 · Sauvegarde", "Game Engine · Player · Catalog"),
        ("3 · Migration", "Bundles EF contrôlés"),
        ("4 · Déploiement", "Images par digest OCI"),
        ("5 · Validation", "Readiness · smoke · E2E critique"),
        ("6 · Décision", "Promouvoir ou rollback applicatif"),
    ]
    y = .85
    for i, (title, detail) in enumerate(stages):
        fill = "#E8F2EE" if i == 5 else "#EEF1F5"
        edge = "#6C9A87" if i == 5 else "#8997AA"
        box(ax, .22, y - .075, .56, .105, title, detail, fill, edge)
        if i < len(stages) - 1:
            arrow(ax, (.5, y - .075), (.5, y - .125))
        y -= .145
    save_figure(fig, "deployment-sequence.png")

    fig, ax = plt.subplots(figsize=(10.5, 6.2))
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    phases = [
        ("M0", "Baseline", "Tests mesurés"),
        ("M1", "CI verte", "80 % par unité"),
        ("M2", "CD", "Images et environnements"),
        ("M3", "Comptes", "Identité et ownership"),
        ("M4", "Vertical slice", "Boucle jouable complète"),
        ("M5", "Contenu", "Parcours bêta"),
        ("M6", "Durcissement", "Ouverture contrôlée"),
    ]
    y = .86
    for i, (code, title, detail) in enumerate(phases):
        color = "#E8F2EE" if i == len(phases) - 1 else "#EEF1F5"
        edge = "#6C9A87" if i == len(phases) - 1 else "#8997AA"
        ax.text(.11, y, code, ha="center", va="center", fontsize=11, fontweight="bold", color="#525C6C")
        box(ax, .20, y - .055, .66, .11, title, detail, color, edge)
        if i < len(phases) - 1:
            arrow(ax, (.53, y - .055), (.53, y - .095))
        y -= .125
    ax.text(.53, .02, "BÊTA JOUABLE, DÉPLOYABLE ET OBSERVABLE", ha="center", va="center",
            fontsize=10, fontweight="bold", color="#1C2027")
    save_figure(fig, "roadmap-beta.png")


def generate(meta: dict) -> Path:
    text = meta["source"].read_text(encoding="utf-8")
    doc = prepare_template(meta)
    configure_body(doc, meta["reference"])
    add_summary(doc, text)
    add_markdown_body(doc, text)
    SOURCE_DIR.mkdir(parents=True, exist_ok=True)
    DELIVERABLE_DIR.mkdir(parents=True, exist_ok=True)
    output = SOURCE_DIR / meta["output"]
    doc.save(output)
    shutil.copy2(output, DELIVERABLE_DIR / output.name)
    return output


if __name__ == "__main__":
    generate_diagrams()
    for document in DOCS:
        print(generate(document))
